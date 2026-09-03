#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Exporta SÓ o dia a dia de outubro/2026 nos formatos que a equipe usa.

Fonte: o Sheet vivo da Mari ("Escala UTI HCB 2026 · V5") — lido, nunca escrito.
    python3 export_dia_a_dia.py ler        # rebaixa o snapshot do Sheet
    python3 export_dia_a_dia.py            # gera tudo a partir do snapshot

Saídas em docs/escala-tools/exports/:
    1. "como está"  — a aba OUT · DIA A DIA: .xlsx .docx .html (→ .pdf)
    2. "como mandam" — a grade do grupo, com o vermelho/amarelo/azul do arquivo
       original que o hospital distribui: .xlsx .docx .html (→ .pdf)
    3. "impresso"   — PDF de verdade, tipografia e grade desenhadas
    4. "painel"     — a leitura mais fácil que consegui: uma página por semana,
       turnos empilhados, tudo visível sem rolar

PDF sai do Chrome headless (--print-to-pdf): texto vetorial, fonte embedada.
"""
import datetime as dt
import json
import os
import subprocess
import sys

AQUI = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, AQUI)
SAIDA = os.path.join(AQUI, "exports")
SNAP = os.path.join(AQUI, "dia_a_dia_out.json")
SHEET = "1SFS8kiq8Qq5ooSyQ22rWlF_z0fJ6jTXLKKmnSvK00_E"
CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
WD = ["segunda", "terça", "quarta", "quinta", "sexta", "sábado", "domingo"]
WD3 = ["seg", "ter", "qua", "qui", "sex", "sáb", "dom"]


# ----------------------------------------------------------------- 0 · o dado
def ler_do_sheet():
    """snapshot do dia a dia a partir do Sheet vivo (só leitura)."""
    import gsuite
    import v4_dados as D
    out = gsuite.ler(SHEET, "OUT!A7:BR80")
    up = {k.upper(): k for k in D.TURNOS}
    lin = {r[0]: r for r in out[1:] if r and r[0]}
    roster = [x[0] for x in D.ROSTER]
    per = {"M-": "manhã", "T-": "tarde", "D-": "dia", "N-": "noite",
           "Tm-": "manhã", "Mt-": "tarde"}

    def cel(ap, i):
        r = lin.get(ap, [])
        v = r[i] if i < len(r) else ""
        v = (v.strip() if isinstance(v, str) else str(v or ""))
        return up.get(v.upper(), v)

    colunas = ([(dt.date(2026, 9, d), c, True) for d, c in ((28, 5), (29, 6), (30, 7))]
               + [(dt.date(2026, 10, d), 8 + d - 1, False) for d in range(1, 32)]
               + [(dt.date(2026, 11, 1), 39, True)])
    dias = []
    for data, i, viz in colunas:
        m, t, n, nt, bhn = [], [], [], [], []
        for ap in roster:
            l = cel(ap, i)
            if not l or l not in D.TURNOS:
                continue
            rot = ap + (" BHP" if l in D.BHP else "")
            cm, ct, cn = D.TURNOS[l][4:7]
            if l == "NT":
                nt.append(ap)
            else:
                if cm:
                    m.append(rot)
                if ct:
                    t.append(rot)
                if cn:
                    n.append(rot)
            if l in D.BHN:
                bhn.append(f"{ap} BHN {per[l]}")
        wd = data.weekday()
        tipo = "útil" if wd < 5 else ("sábado" if wd == 5 else "domingo")
        mn = D.MINIMOS[tipo] if data.month >= 10 else D.MINIMOS_ANTIGOS[tipo]
        fer = D.FERIADOS_2026.get((data.month, data.day))
        cob = [len(m), len(t), len(n)]
        dias.append(dict(
            data=f"{data:%Y-%m-%d}", label=f"{data:%d/%m}", dia=data.day, mes=data.month,
            dow=WD[wd], dow3=WD3[wd], fds=wd >= 5, vizinho=viz,
            feriado=(fer[0] if fer else None), sigla=(fer[1] if fer else None),
            manha=m, tarde=t, noite=n, noitinha=nt, bhn=bhn,
            cob=cob, min=list(mn), falta=[max(0, mn[k] - cob[k]) for k in range(3)]))
    meta = dict(gerado=f"{dt.datetime.now():%d/%m/%Y %H:%M}", fonte="Sheet · V5 (leitura)")
    json.dump({"meta": meta, "dias": dias}, open(SNAP, "w"), ensure_ascii=False, indent=1)
    return dias, meta


def carregar():
    d = json.load(open(SNAP))
    return d["dias"], d["meta"]


def semanas(dias):
    """agrupa em semanas civis (segunda a domingo), como a planilha conta."""
    blocos, atual = [], []
    for d in dias:
        if atual and d["dow3"] == "seg":
            blocos.append(atual)
            atual = []
        atual.append(d)
    if atual:
        blocos.append(atual)
    return blocos


def pdf(html_path, pdf_path, paisagem=False):
    """Chrome headless: PDF vetorial com as fontes embedadas."""
    cmd = [CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
           "--virtual-time-budget=6000", f"--print-to-pdf={pdf_path}",
           "file://" + os.path.abspath(html_path)]
    subprocess.run(cmd, capture_output=True, timeout=120)
    return os.path.exists(pdf_path)


# ============================================================ 1 · COMO ESTÁ
# A aba OUT · DIA A DIA como ela é hoje: uma linha por dia, os nomes de cada
# turno na célula, os BHN em coluna própria, a cobertura no fim.
INK, INK2, INK3 = "3A2E2A", "6B5C56", "9A8A82"
CREME, CREME2, LINE = "FFFAF3", "FAF3E8", "EBE8E5"
LAV, LAVI, LAVS = "A299CB", "5A4E8C", "ECEAF4"
CORAL, CORALI, CORALS = "E7A59C", "C77264", "FBE9E5"
COLS1 = [("Dia", 7), ("", 9), ("Manhã", 62), ("Tarde", 50), ("Noite", 42),
         ("BHN · dispensa", 26), ("Cobertura", 14)]


def _linhas_como_esta(dias):
    for d in dias:
        cob = (f"M {d['cob'][0]}/{d['min'][0]} · T {d['cob'][1]}/{d['min'][1]} · "
               f"N {d['cob'][2]}/{d['min'][2]}")
        if sum(d["falta"]):
            cob += "  FALTA " + " ".join(f"{r}{f}" for r, f in
                                        zip("MTN", d["falta"]) if f)
        rot = d["label"] + ("  ·  " + d["feriado"] if d["feriado"] else "")
        yield d, rot, cob


def f1_xlsx(dias, meta):
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    fino = Side(style="thin", color=LINE)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "OUT · DIA A DIA"
    ws.sheet_view.showGridLines = False
    t = ws.cell(row=1, column=1, value="Outubro de 2026 · dia a dia")
    t.font = Font(name="Fraunces", bold=True, size=15, color=INK)
    s = ws.cell(row=2, column=1, value=f"Escala UTI HCB · {meta['fonte']} · "
                f"gerado em {meta['gerado']} · a semana fecha no domingo, por isso "
                "aparecem 28–30/09 e 01/11")
    s.font = Font(name="Nunito", size=9, italic=True, color=INK3)
    for j, (h, w) in enumerate(COLS1, start=1):
        c = ws.cell(row=4, column=j, value=h)
        c.font = Font(name="Nunito", bold=True, size=9, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=LAVI)
        c.alignment = Alignment(horizontal="left", vertical="center")
        ws.column_dimensions[get_column_letter(j)].width = w
    r = 5
    for d, rot, cob in _linhas_como_esta(dias):
        vals = [rot, d["dow"], ", ".join(d["manha"]) or "—", ", ".join(d["tarde"]) or "—",
                ", ".join(d["noite"]) or "—", ", ".join(d["bhn"]), cob]
        for j, v in enumerate(vals, start=1):
            c = ws.cell(row=r, column=j, value=v)
            c.font = Font(name="Nunito", size=8,
                          bold=(j == 1 and not d["vizinho"]),
                          italic=d["vizinho"],
                          color=(INK3 if d["vizinho"] else
                                 (CORALI if (j == 7 and sum(d["falta"])) else INK)))
            c.alignment = Alignment(wrap_text=(j >= 3), vertical="top")
            c.border = Border(left=fino, right=fino, top=fino, bottom=fino)
            if d["feriado"]:
                c.fill = PatternFill("solid", fgColor=CORALS)
            elif d["vizinho"]:
                c.fill = PatternFill("solid", fgColor=LINE)
            elif d["fds"]:
                c.fill = PatternFill("solid", fgColor=LAVS)
            elif r % 2:
                c.fill = PatternFill("solid", fgColor=CREME2)
        import math
        maior = max(len(", ".join(d[k])) / w for k, (_h, w) in
                    zip(("manha", "tarde", "noite"), COLS1[2:5]))
        ws.row_dimensions[r].height = max(15, math.ceil(maior) * 11 + 6)
        r += 1
    ws.freeze_panes = "C5"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    p = os.path.join(SAIDA, "1 · dia a dia · como está.xlsx")
    wb.save(p)
    return p


def f1_docx(dias, meta):
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.2))
    h = doc.add_paragraph()
    run = h.add_run("Outubro de 2026 · dia a dia")
    run.font.name, run.font.size, run.bold = "Georgia", Pt(16), True
    sub = doc.add_paragraph()
    r2 = sub.add_run(f"Escala UTI HCB · {meta['fonte']} · gerado em {meta['gerado']}")
    r2.font.name, r2.font.size, r2.italic = "Calibri", Pt(8)  , True
    r2.font.color.rgb = RGBColor(0x6B, 0x5C, 0x56)
    tb = doc.add_table(rows=1, cols=6)
    tb.style = "Table Grid"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    larg = [Cm(2.0), Cm(8.6), Cm(6.8), Cm(5.6), Cm(3.4), Cm(2.6)]
    for j, txt in enumerate(("Dia", "Manhã", "Tarde", "Noite", "BHN · dispensa", "Cobertura")):
        cel = tb.rows[0].cells[j]
        cel.text = ""
        rr = cel.paragraphs[0].add_run(txt)
        rr.bold = True
        rr.font.size, rr.font.name = Pt(8), "Calibri"
        cel.width = larg[j]
    for d, rot, cob in _linhas_como_esta(dias):
        row = tb.add_row()
        vals = [f"{rot}\n{d['dow3']}", ", ".join(d["manha"]) or "—",
                ", ".join(d["tarde"]) or "—", ", ".join(d["noite"]) or "—",
                ", ".join(d["bhn"]), cob]
        for j, v in enumerate(vals):
            cel = row.cells[j]
            cel.text = ""
            cel.width = larg[j]
            rr = cel.paragraphs[0].add_run(v)
            rr.font.size, rr.font.name = Pt(7), "Calibri"
            rr.bold = (j == 0 and not d["vizinho"])
            rr.italic = d["vizinho"]
            if d["vizinho"]:
                rr.font.color.rgb = RGBColor(0x9A, 0x8A, 0x82)
            elif j == 5 and sum(d["falta"]):
                rr.font.color.rgb = RGBColor(0xC7, 0x72, 0x64)
    p = os.path.join(SAIDA, "1 · dia a dia · como está.docx")
    doc.save(p)
    return p


# ========================================================= 2 · COMO MANDAM
# A grade que o hospital distribui, com o vermelho/amarelo/azul do arquivo
# original: título do mês em vermelho, faixa da semana em azul-marinho,
# cabeçalho dos turnos em amarelo com letra azul, um nome por linha.
# Cores lidas do "Escala setembro - UTI HCB.xlsx": FF0000 · FFFF00 · 0000FF ·
# 000080 · 002060 · D9E1F2 (o azul claro dos destaques).
VERM, AMAR, AZUL, MARINHO, AZ_NOME, AZ_CLARO = ("FF0000", "FFFF00", "0000FF",
                                                "000080", "002060", "D9E1F2")
COLS2 = [("Coord", 14), ("", 12), ("Dia", 6), ("Manhã", 24), ("Tarde", 22),
         ("Noite", 20), ("Noitinha", 14)]
ORD = ("1ª", "2ª", "3ª", "4ª", "5ª", "6ª")


def _blocos_como_mandam(dias):
    """(semana, dia, linhas) — cada linha é (manhã, tarde, noite, noitinha)."""
    for k, bloco in enumerate(semanas(dias)):
        linhas_semana = []
        for d in bloco:
            n = max(len(d["manha"]), len(d["tarde"]), len(d["noite"]),
                    len(d["noitinha"]), 1)
            # os BHN entram no fim do dia, como na grade oficial
            extra = [(x, "", "", "") for x in d["bhn"]]
            corpo = [(d["manha"][i] if i < len(d["manha"]) else "",
                      d["tarde"][i] if i < len(d["tarde"]) else "",
                      d["noite"][i] if i < len(d["noite"]) else "",
                      d["noitinha"][i] if i < len(d["noitinha"]) else "")
                     for i in range(n)]
            linhas_semana.append((d, corpo + extra))
        yield ORD[k] + " Semana", linhas_semana


def f2_xlsx(dias, meta):
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    fino = Side(style="thin", color="000000")
    box = Border(left=fino, right=fino, top=fino, bottom=fino)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Outubro 2026"
    for j, (_h, w) in enumerate(COLS2, start=1):
        ws.column_dimensions[get_column_letter(j)].width = w
    r = 1
    tit = ws.cell(row=r, column=1, value="OUTUBRO")
    tit.font = Font(name="Calibri", bold=True, size=18, color="FFFFFF")
    tit.fill = PatternFill("solid", fgColor=VERM)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    r += 1
    for nome_sem, linhas_semana in _blocos_como_mandam(dias):
        cs = ws.cell(row=r, column=2, value=nome_sem)
        cs.font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
        cs.fill = PatternFill("solid", fgColor=MARINHO)
        r += 1
        for j, (h, _w) in enumerate(COLS2, start=1):
            c = ws.cell(row=r, column=j, value=h or None)
            c.font = Font(name="Calibri Light", bold=True, size=12, color=AZUL)
            c.fill = PatternFill("solid", fgColor=AMAR)
            c.alignment = Alignment(horizontal="center")
            c.border = box
        r += 1
        for d, corpo in linhas_semana:
            r0 = r
            for i, (m, t, n, nt) in enumerate(corpo):
                for j, v in enumerate((m, t, n, nt), start=4):
                    c = ws.cell(row=r, column=j, value=v or None)
                    anot = any(x in v for x in ("BHP", "BHN", "CEP", "CP", "CRO"))
                    c.font = Font(name="Calibri", size=11, bold=anot,
                                  color=(AZ_NOME if anot else "000000"))
                    c.border = box
                    if d["feriado"]:
                        c.fill = PatternFill("solid", fgColor="FFC7CE")
                    elif d["fds"]:
                        c.fill = PatternFill("solid", fgColor=AZ_CLARO)
                for j in (1, 2, 3):
                    ws.cell(row=r, column=j).border = box
                r += 1
            cd = ws.cell(row=r0, column=2,
                         value=d["dow"].capitalize() + (" (set)" if d["mes"] == 9
                                                        else " (nov)" if d["mes"] == 11 else ""))
            cd.font = Font(name="Calibri", bold=True, size=11)
            cd.alignment = Alignment(horizontal="center", vertical="center")
            cn = ws.cell(row=r0, column=3, value=d["label"] if d["mes"] != 10 else d["dia"])
            cn.font = Font(name="Calibri", bold=True, size=11,
                           color=(VERM if d["feriado"] else "000000"))
            cn.alignment = Alignment(horizontal="center", vertical="center")
            if r > r0 + 1:
                ws.merge_cells(start_row=r0, start_column=2, end_row=r - 1, end_column=2)
                ws.merge_cells(start_row=r0, start_column=3, end_row=r - 1, end_column=3)
            if d["feriado"]:
                cf = ws.cell(row=r0, column=1, value=d["feriado"].upper())
                cf.font = Font(name="Calibri", bold=True, size=10, color=VERM)
        r += 1
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    p = os.path.join(SAIDA, "2 · dia a dia · como mandam.xlsx")
    wb.save(p)
    return p


def f2_docx(dias, meta):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement
    def sombrear(cel, hexa):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexa)
        cel._tc.get_or_add_tcPr().append(el)
    doc = Document()
    sec = doc.sections[0]
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.4))
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("OUTUBRO")
    r0.font.name, r0.font.size, r0.bold = "Calibri", Pt(18), True
    r0.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    larg = [Cm(2.6), Cm(1.4), Cm(4.6), Cm(4.2), Cm(3.8), Cm(2.4)]
    for nome_sem, linhas_semana in _blocos_como_mandam(dias):
        ps = doc.add_paragraph()
        rs = ps.add_run(nome_sem)
        rs.font.name, rs.font.size, rs.bold = "Calibri", Pt(11), True
        rs.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
        tb = doc.add_table(rows=1, cols=6)
        tb.style = "Table Grid"
        for j, txt in enumerate(("Dia", "", "Manhã", "Tarde", "Noite", "Noitinha")):
            cel = tb.rows[0].cells[j]
            cel.text = ""
            cel.width = larg[j]
            sombrear(cel, AMAR)
            rr = cel.paragraphs[0].add_run(txt)
            rr.font.name, rr.font.size, rr.bold = "Calibri Light", Pt(10), True
            rr.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        for d, corpo in linhas_semana:
            row = tb.add_row()
            row.cells[0].text = ""
            rd = row.cells[0].paragraphs[0].add_run(
                f"{d['dow'].capitalize()}\n{d['label']}"
                + (f"\n{d['feriado'].upper()}" if d["feriado"] else ""))
            rd.font.name, rd.font.size, rd.bold = "Calibri", Pt(9), True
            if d["feriado"]:
                rd.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            row.cells[1].text = ""
            for j, campo in enumerate(("manha", "tarde", "noite", "noitinha"), start=2):
                cel = row.cells[j]
                cel.text = ""
                cel.width = larg[j]
                if d["feriado"]:
                    sombrear(cel, "FFC7CE")
                elif d["fds"]:
                    sombrear(cel, AZ_CLARO)
                nomes = d[campo] + (d["bhn"] if campo == "manha" else [])
                par = cel.paragraphs[0]
                for i, nm in enumerate(nomes):
                    rr = par.add_run(("\n" if i else "") + nm)
                    anot = any(x in nm for x in ("BHP", "BHN", "CEP", "CP", "CRO"))
                    rr.font.name, rr.font.size, rr.bold = "Calibri", Pt(9), anot
                    if anot:
                        rr.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    p = os.path.join(SAIDA, "2 · dia a dia · como mandam.docx")
    doc.save(p)
    return p


# ------------------------------------------------------------- html + pdf
def _esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _nomes_html(lista, classe=""):
    saida = []
    for n in lista:
        cl = classe + (" bhp" if n.endswith("BHP") else "")
        saida.append(f'<span class="n {cl}">{_esc(n)}</span>')
    return "".join(saida) or '<span class="vazio">—</span>'


CABECA = """<meta charset="utf-8"><title>{titulo}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,600;9..144,700&family=Nunito:wght@400;600;700&display=swap" rel="stylesheet">"""


def f1_html(dias, meta):
    linhas = []
    for d, rot, cob in _linhas_como_esta(dias):
        cls = " ".join(filter(None, ["viz" if d["vizinho"] else "",
                                     "fds" if d["fds"] else "",
                                     "fer" if d["feriado"] else ""]))
        linhas.append(f"""<tr class="{cls}">
<td class="dia"><b>{_esc(d['label'])}</b><br><span class="dow">{_esc(d['dow'])}</span>
{f'<br><span class="fer-nome">{_esc(d["feriado"])}</span>' if d['feriado'] else ''}</td>
<td>{_nomes_html(d['manha'])}</td><td>{_nomes_html(d['tarde'])}</td>
<td>{_nomes_html(d['noite'])}</td><td class="bhn">{_nomes_html(d['bhn']) if d['bhn'] else ''}</td>
<td class="cob {'falta' if sum(d['falta']) else ''}">{_esc(cob)}</td></tr>""")
    html = f"""{CABECA.format(titulo='Outubro 2026 · dia a dia')}
<style>
@page {{ size: A4 landscape; margin: 10mm; }}
* {{ box-sizing: border-box; }}
body {{ font-family: Nunito, system-ui, sans-serif; color: #3A2E2A; background: #FFFAF3;
       margin: 0; font-size: 8pt; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
h1 {{ font-family: Fraunces, Georgia, serif; font-size: 17pt; margin: 0 0 2pt; font-weight: 600; }}
.sub {{ color: #6B5C56; font-size: 8pt; font-style: italic; margin-bottom: 7pt; }}
table {{ border-collapse: collapse; width: 100%; }}
th {{ background: #5A4E8C; color: #fff; font-size: 8pt; text-align: left; padding: 4pt 5pt;
      font-weight: 700; letter-spacing: .01em; }}
td {{ border: .4pt solid #EBE8E5; padding: 3pt 5pt; vertical-align: top; }}
tr:nth-child(even) td {{ background: #FAF3E8; }}
tr.fds td {{ background: #ECEAF4; }}
tr.fer td {{ background: #FBE9E5; }}
tr.viz td {{ background: #F2F0ED; color: #9A8A82; font-style: italic; }}
.dia {{ white-space: nowrap; font-variant-numeric: tabular-nums; }}
.dow {{ color: #6B5C56; font-size: 7pt; }}
.fer-nome {{ color: #C77264; font-size: 6.5pt; font-weight: 700; }}
.n:not(:last-child)::after {{ content: ", "; color: #9A8A82; }}
.n.bhp {{ color: #C77264; font-weight: 700; }}
.bhn {{ color: #6B5C56; font-style: italic; }}
.vazio {{ color: #9A8A82; }}
.cob {{ white-space: nowrap; font-variant-numeric: tabular-nums; color: #6B5C56; }}
.cob.falta {{ color: #C77264; font-weight: 700; }}
</style>
<h1>Outubro de 2026 · dia a dia</h1>
<div class="sub">Escala UTI HCB · {_esc(meta['fonte'])} · gerado em {_esc(meta['gerado'])} ·
a semana fecha no domingo, por isso aparecem 28–30/09 e 01/11</div>
<table><thead><tr><th>Dia</th><th>Manhã</th><th>Tarde</th><th>Noite</th>
<th>BHN · dispensa</th><th>Cobertura</th></tr></thead><tbody>
{''.join(linhas)}</tbody></table>"""
    p = os.path.join(SAIDA, "1 · dia a dia · como está.html")
    open(p, "w").write(html)
    return p


def f2_html(dias, meta):
    blocos = []
    for nome_sem, linhas_semana in _blocos_como_mandam(dias):
        linhas = []
        for d, _corpo in linhas_semana:
            cls = "fer" if d["feriado"] else ("fds" if d["fds"] else "")
            def col(campo):
                nomes = d[campo]
                return "<br>".join(
                    f'<span class="{"anot" if any(x in n for x in ("BHP","BHN","CEP","CP","CRO")) else ""}">{_esc(n)}</span>'
                    for n in nomes) or "&nbsp;"
            bhn = "<br>".join(f'<span class="anot">{_esc(x)}</span>' for x in d["bhn"])
            linhas.append(f"""<tr class="{cls}">
<td class="coord">{_esc(d['feriado'].upper()) if d['feriado'] else ''}</td>
<td class="dow">{_esc(d['dow'].capitalize())}</td>
<td class="num">{d['dia'] if d['mes']==10 else _esc(d['label'])}</td>
<td>{col('manha')}{('<br>'+bhn) if bhn else ''}</td><td>{col('tarde')}</td>
<td>{col('noite')}</td><td>{col('noitinha')}</td></tr>""")
        blocos.append(f"""<div class="sem">{_esc(nome_sem)}</div>
<table><thead><tr><th>Coord</th><th></th><th>Dia</th><th>Manhã</th><th>Tarde</th>
<th>Noite</th><th>Noitinha</th></tr></thead><tbody>{''.join(linhas)}</tbody></table>""")
    html = f"""<meta charset="utf-8"><title>OUTUBRO</title>
<style>
@page {{ size: A4; margin: 12mm; }}
body {{ font-family: Calibri, "Segoe UI", sans-serif; font-size: 10pt; color: #000;
        background: #fff; margin: 0; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
h1 {{ background: #FF0000; color: #fff; font-size: 18pt; margin: 0 0 6pt; padding: 2pt 6pt;
      font-weight: 700; }}
.sem {{ background: #000080; color: #fff; font-weight: 700; padding: 1pt 6pt;
        margin: 10pt 0 0; display: inline-block; }}
table {{ border-collapse: collapse; width: 100%; margin-bottom: 4pt; }}
th {{ background: #FFFF00; color: #0000FF; border: .5pt solid #000; font-family: "Calibri Light", Calibri;
      font-size: 11pt; font-weight: 700; padding: 1pt 3pt; }}
td {{ border: .5pt solid #000; padding: 1pt 3pt; vertical-align: top; line-height: 1.25; }}
tr.fds td {{ background: #D9E1F2; }}
tr.fer td {{ background: #FFC7CE; }}
.dow, .num {{ font-weight: 700; text-align: center; white-space: nowrap; }}
.coord {{ color: #FF0000; font-weight: 700; font-size: 8pt; }}
.anot {{ color: #002060; font-weight: 700; }}
</style>
<h1>OUTUBRO</h1>
{''.join(blocos)}"""
    p = os.path.join(SAIDA, "2 · dia a dia · como mandam.html")
    open(p, "w").write(html)
    return p


# ======================================================== 3 · IMPRESSO
# Uma página por semana civil, paisagem. O dia é a coluna; manhã, tarde e noite
# são faixas. Quem lê procura um nome numa coluna estreita, então o nome é a
# unidade tipográfica: uma linha, sem vírgula, alinhado à esquerda.
def f3_html(dias, meta):
    paginas = []
    for k, bloco in enumerate(semanas(dias)):
        ini, fim = bloco[0], bloco[-1]
        cabecas, faixas = [], {"manha": [], "tarde": [], "noite": []}
        pes = []
        for d in bloco:
            marca = ("feriado" if d["feriado"] else
                     ("fds" if d["fds"] else ("viz" if d["vizinho"] else "")))
            cabecas.append(f"""<th class="{marca}">
<span class="num">{d['dia']:02d}</span><span class="mes">/{d['mes']:02d}</span>
<span class="dow">{_esc(d['dow'])}</span>
{f'<span class="fer">{_esc(d["feriado"])}</span>' if d['feriado'] else ''}</th>""")
            for campo in ("manha", "tarde", "noite"):
                itens = "".join(
                    f'<li class="{"bhp" if n.endswith(" BHP") else ""}">{_esc(n.replace(" BHP",""))}'
                    f'{"<i>BHP</i>" if n.endswith(" BHP") else ""}</li>'
                    for n in d[campo])
                faixas[campo].append(f'<td class="{marca}"><ul>{itens or "<li class=vazio>—</li>"}</ul></td>')
            falta = sum(d["falta"])
            pes.append(f"""<td class="{marca}">
<span class="cob {'falta' if falta else ''}">{d['cob'][0]}·{d['cob'][1]}·{d['cob'][2]}</span>
{f'<span class="buraco">faltam {falta}</span>' if falta else ''}
{''.join(f'<span class="bhn">{_esc(x)}</span>' for x in d['bhn'])}</td>""")
        paginas.append(f"""<section>
<header><h2>{ORD[k]} semana</h2>
<span class="periodo">{_esc(ini['label'])} a {_esc(fim['label'])}</span></header>
<table><thead><tr><th class="canto"></th>{''.join(cabecas)}</tr></thead><tbody>
<tr><th class="turno">manhã<span>07–13h</span></th>{''.join(faixas['manha'])}</tr>
<tr><th class="turno">tarde<span>13–19h</span></th>{''.join(faixas['tarde'])}</tr>
<tr><th class="turno">noite<span>19–07h</span></th>{''.join(faixas['noite'])}</tr>
<tr class="rodape"><th class="turno">no dia<span>M·T·N</span></th>{''.join(pes)}</tr>
</tbody></table></section>""")
    html = f"""{CABECA.format(titulo='Escala UTI HCB · outubro de 2026')}
<style>
@page {{ size: A4 landscape; margin: 9mm 10mm; }}
:root {{ --tinta:#3A2E2A; --tinta2:#6B5C56; --tinta3:#9A8A82; --creme:#FFFAF3;
         --creme2:#FAF3E8; --linha:#E4DFD9; --lav:#5A4E8C; --lavs:#EFEDF6;
         --coral:#B65B4B; --corals:#FBE9E5; --areia:#E8C79A; }}
* {{ box-sizing: border-box; }}
body {{ margin:0; background: var(--creme); color: var(--tinta);
        font-family: Nunito, system-ui, sans-serif; font-size: 7.6pt; line-height: 1.28;
        -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
::selection {{ background: var(--lavs); color: var(--lav); }}
section {{ page-break-after: always; padding-top: 1mm; }}
section:last-child {{ page-break-after: auto; }}
header {{ display:flex; align-items: baseline; gap: 4mm; border-bottom: .8pt solid var(--tinta);
          padding-bottom: 1.6mm; margin-bottom: 2.4mm; }}
h2 {{ font-family: Fraunces, Georgia, serif; font-weight: 600; font-size: 15pt; margin:0;
      letter-spacing: -0.01em; }}
.periodo {{ font-size: 8.4pt; color: var(--tinta2); font-variant-numeric: tabular-nums; }}
header::after {{ content: "Escala UTI HCB · outubro de 2026"; margin-left: auto;
                 font-size: 7pt; color: var(--tinta3); letter-spacing: .04em;
                 text-transform: uppercase; }}
table {{ border-collapse: collapse; width: 100%; table-layout: fixed; }}
th, td {{ border: .4pt solid var(--linha); vertical-align: top; padding: 1.2mm 1.4mm; }}
thead th {{ border-bottom: .8pt solid var(--tinta); text-align: left; padding-bottom: 1.4mm; }}
.canto, .turno {{ width: 15mm; background: var(--creme2); }}
.turno {{ font-family: Fraunces, Georgia, serif; font-size: 8.6pt; font-weight: 600;
          color: var(--lav); text-align: left; }}
.turno span {{ display:block; font-family: Nunito, sans-serif; font-size: 6.4pt;
               font-weight: 400; color: var(--tinta3); font-variant-numeric: tabular-nums; }}
.num {{ font-family: Fraunces, Georgia, serif; font-size: 15pt; font-weight: 600;
        font-variant-numeric: tabular-nums; letter-spacing: -0.02em; }}
.mes {{ font-size: 7pt; color: var(--tinta3); font-variant-numeric: tabular-nums; }}
.dow {{ display:block; font-size: 7pt; color: var(--tinta2); text-transform: lowercase; }}
.fer {{ display:block; font-size: 6.2pt; font-weight:700; color: var(--coral);
        text-transform: uppercase; letter-spacing: .03em; }}
thead th.fds .num, thead th.fds .dow {{ color: var(--lav); }}
thead th.feriado .num {{ color: var(--coral); }}
td.fds {{ background: var(--lavs); }}
td.feriado {{ background: var(--corals); }}
td.viz {{ background: #F3F1EE; color: var(--tinta3); }}
ul {{ margin:0; padding:0; list-style:none; }}
li {{ white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
li.bhp {{ color: var(--coral); font-weight: 700; }}
li.bhp i {{ font-style: normal; font-size: 5.6pt; letter-spacing:.04em; margin-left: .6mm;
            vertical-align: .3mm; }}
li.vazio {{ color: var(--tinta3); }}
tr.rodape td, tr.rodape th {{ border-top: .8pt solid var(--tinta); background: var(--creme2); }}
tr.rodape td.fds {{ background: var(--lavs); }}
.cob {{ font-family: Fraunces, Georgia, serif; font-size: 9pt; font-weight: 600;
        font-variant-numeric: tabular-nums; color: var(--tinta2); }}
.cob.falta {{ color: var(--coral); }}
.buraco, .bhn {{ display:block; font-size: 6.2pt; color: var(--coral); font-weight: 700; }}
.bhn {{ color: var(--tinta2); font-weight: 400; font-style: italic; }}
</style>
{''.join(paginas)}"""
    p = os.path.join(SAIDA, "3 · dia a dia · impresso.html")
    open(p, "w").write(html)
    return p
